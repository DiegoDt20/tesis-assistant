"""Servicio de Exportación — genera archivos .docx desde un proyecto."""
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.models.proyecto import Proyecto
from app.models.plantilla import Plantilla


def exportar_a_docx(
    proyecto: Proyecto,
    plantilla: Plantilla,
) -> bytes:
    """
    Genera un archivo .docx con el contenido del proyecto.
    Devuelve los bytes del archivo para enviarlo como descarga.
    """
    doc = Document()

    # ---------------------------------------------------------------------------
    # Portada
    # ---------------------------------------------------------------------------
    titulo_para = doc.add_paragraph()
    titulo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo_para.add_run(proyecto.titulo)
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph()  # espacio

    # Información básica
    info = doc.add_paragraph()
    info.add_run(f"Plantilla: {plantilla.titulo}").italic = True
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ---------------------------------------------------------------------------
    # Secciones del documento
    # ---------------------------------------------------------------------------
    respuestas = proyecto.respuestas_json or {}

    for seccion in plantilla.secciones_json:
        seccion_id = seccion.get("id", "")
        seccion_titulo = seccion.get("titulo", seccion_id)

        # Título de la sección
        heading = doc.add_heading(seccion_titulo, level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        # Contenido de la sección
        contenido = respuestas.get(seccion_id, "")
        if contenido:
            doc.add_paragraph(contenido)
        else:
            # Sección sin contenido
            parrafo = doc.add_paragraph()
            run = parrafo.add_run("[Sección pendiente de completar]")
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # Procesar sub-secciones si existen
        for hijo in seccion.get("hijos", []):
            hijo_id = hijo.get("id", "")
            hijo_titulo = hijo.get("titulo", hijo_id)

            heading2 = doc.add_heading(hijo_titulo, level=2)
            contenido_hijo = respuestas.get(hijo_id, "")
            if contenido_hijo:
                doc.add_paragraph(contenido_hijo)
            else:
                parrafo = doc.add_paragraph()
                run = parrafo.add_run("[Sección pendiente de completar]")
                run.italic = True
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        doc.add_paragraph()  # espacio entre secciones

    # ---------------------------------------------------------------------------
    # Guardar en memoria y devolver bytes
    # ---------------------------------------------------------------------------
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()